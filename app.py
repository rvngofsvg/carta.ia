
import io
import re
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# =========================================================
# CONFIGURACIÓN PRINCIPAL
# =========================================================
APP_TITLE = "Carta Alérgenos"
DEFAULT_LEGEND_PATH = Path("public/leyenda_alergenos.png")

# Esta plantilla es la que pidió la clienta:
# texto + precios, SIN iconos por plato, manteniendo la leyenda inferior.
TEMPLATE_WORD_TEXTO_LIMPIO = "Word texto limpio"


DEFAULT_MENU_TEXT = """# DE PICOTEO
Tortilla de patatas | 3,00 € | gluten
Magra con tomate | 6,00 € |
Calamares con tomate | 6,00 € | pescado
Ensaladilla rusa | 5,00 € |
Ensaladilla de marisco | 6,00 € |
Marinera/O | 3,00 € |
Bicicleta | 2,50 € |
Morcilla de Burgos | 12,00 € |
Tortitas de camarón | 3,20 € |
Caballitos | 3,50 € |
Alitas de pollo clásicas (6 unidades) | 6,50 € |
Alitas de pollo búfalo (4 unidades) | 6,50 € |
Nachos Barnuevo | 15,00 € |
Salchicha seca con almendras Marcona | 12,00 € |
Tosta de tomate con sardina ahumada y salsa Noruega | 6,50 € |

# PLATOS
Jamón Ibérico y queso parmesano | 28,00 € |
Jamón Ibérico con almendras | 20,00 € |
Queso parmesano con almendras | 10,00 € |
Costillar barbacoa con patatas | 22,00 € |
Berenjena boloñesa gratinada con patatas chips | 12,00 € |
Tabla de pates y quesos | 22,50 € |

# ENSALADAS (15,00 €)
Tomate con bonito y aceitunas | 15,00 € |
Ensalada tropical | 15,00 € |
Ensalada César | 15,00 € |

# PIZZAS MASA FINA (12,00 €)
Jamón y queso | 12,00 € |
Vegana | 12,00 € |
Con 4 quesos | 12,00 € |
A la carbonara | 12,00 € |
Al barbacoa | 12,00 € |
Con atún | 12,00 € |
"""


# =========================================================
# UTILIDADES DOCX
# =========================================================
def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_no_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)


def set_cell_margins(cell, top=40, start=0, bottom=40, end=0) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=10, bold=False, color="FFFFFF", italic=False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_black_paragraph(document, text="", size=10, bold=False, italic=False, align=None, space_after=0):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic)
    return p


def set_document_page_background(document, color="000000") -> None:
    """
    Fondo visual negro mediante tabla contenedora.
    Word no maneja de forma universal el color de página vía python-docx,
    así que el contenido se crea en tablas con celdas negras.
    """
    section = document.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)


def set_document_defaults(document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(255, 255, 255)


def make_full_width_black_table(document, rows=1, cols=1):
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_no_borders(table)
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "000000")
            set_cell_margins(cell, top=20, start=20, bottom=20, end=20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_text_to_cell(cell, text, size=10, bold=False, italic=False, color="FFFFFF", align=None, space_after=0):
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]

    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)

    # Limpia contenido previo del párrafo.
    for run in p.runs:
        run.text = ""

    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def find_legend_image() -> Path | None:
    """
    Mantiene la imagen de la carpeta public.
    No la copia, no la borra y no la reemplaza.
    Busca primero public/leyenda_alergenos.png.
    Si el nombre fuera distinto, intenta localizar una imagen que contenga
    'leyenda' o 'alergen' en public.
    """
    if DEFAULT_LEGEND_PATH.exists():
        return DEFAULT_LEGEND_PATH

    public_dir = Path("public")
    if not public_dir.exists():
        return None

    valid_ext = {".png", ".jpg", ".jpeg", ".webp"}
    candidates = []
    for path in public_dir.rglob("*"):
        if path.suffix.lower() in valid_ext:
            name = path.name.lower()
            if "leyenda" in name or "alergen" in name or "alérgen" in name:
                candidates.append(path)

    return candidates[0] if candidates else None


# =========================================================
# PARSEO DE MENÚ
# =========================================================
def normalize_price(price: str) -> str:
    price = price.strip()
    if not price:
        return ""
    price = price.replace("€", "").strip()
    return f"{price} €"


def parse_menu_text(raw_text: str):
    """
    Formato admitido:
    # CATEGORÍA
    Plato | Precio | alergenos opcionales

    En Word texto limpio se ignora la tercera columna de alérgenos.
    Se conserva por compatibilidad con otras plantillas.
    """
    sections = []
    current = None

    for raw_line in raw_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("#"):
            title = line.lstrip("#").strip()
            current = {"title": title, "items": []}
            sections.append(current)
            continue

        if current is None:
            current = {"title": "MENÚ", "items": []}
            sections.append(current)

        parts = [p.strip() for p in line.split("|")]
        name = parts[0] if len(parts) >= 1 else ""
        price = normalize_price(parts[1]) if len(parts) >= 2 else ""
        allergens = parts[2] if len(parts) >= 3 else ""

        if name:
            current["items"].append({
                "name": name,
                "price": price,
                "allergens": allergens,
            })

    return sections


def clean_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[^\w\s-]", "", name, flags=re.UNICODE)
    name = re.sub(r"[\s_]+", "_", name)
    return name or "carta"


# =========================================================
# GENERADOR WORD TEXTO LIMPIO
# =========================================================
def build_word_texto_limpio_docx(
    restaurant_name: str,
    sections: list[dict],
    include_footer_legend: bool = True,
    include_footer_text: bool = True,
    legend_width_cm: float = 17.2,
) -> bytes:
    """
    Genera el Word editable limpio:
    - SIN iconos de alérgenos por plato.
    - Plato a la izquierda.
    - Precio alineado a la derecha.
    - Mantiene la imagen de leyenda inferior desde /public.
    """
    document = Document()
    set_document_defaults(document)
    set_document_page_background(document)

    # Título principal
    header_table = make_full_width_black_table(document, 1, 1)
    header_cell = header_table.cell(0, 0)
    add_text_to_cell(header_cell, restaurant_name.upper() if restaurant_name else "MENÚ", size=24, bold=True, space_after=4)

    # Contenido de secciones
    for section in sections:
        title = section.get("title", "").strip()
        items = section.get("items", [])

        if not title and not items:
            continue

        section_table = make_full_width_black_table(document, 1, 1)
        section_cell = section_table.cell(0, 0)
        add_text_to_cell(section_cell, title.upper(), size=16, bold=True, space_after=2)

        # Tabla interna: 2 columnas, nombre y precio.
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_no_borders(table)

        for item in items:
            row = table.add_row()
            name_cell = row.cells[0]
            price_cell = row.cells[1]

            for cell in row.cells:
                set_cell_shading(cell, "000000")
                set_cell_margins(cell, top=25, start=0, bottom=25, end=0)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                remove_cell_borders(cell)

            # Ancho aproximado para que el precio quede más tirado a la derecha.
            name_cell.width = Cm(13.6)
            price_cell.width = Cm(3.0)

            add_text_to_cell(
                name_cell,
                item.get("name", ""),
                size=9.5,
                bold=True,
                color="FFFFFF",
                align=WD_ALIGN_PARAGRAPH.LEFT
            )
            add_text_to_cell(
                price_cell,
                item.get("price", ""),
                size=9.5,
                bold=True,
                color="FFFFFF",
                align=WD_ALIGN_PARAGRAPH.RIGHT
            )

        # Espacio suave entre categorías.
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)

    # Leyenda inferior: misma imagen dentro de public.
    legend_path = find_legend_image()
    if include_footer_legend and legend_path and legend_path.exists():
        legend_table = make_full_width_black_table(document, 1, 1)
        legend_cell = legend_table.cell(0, 0)
        p = legend_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(str(legend_path), width=Cm(legend_width_cm))

    elif include_footer_legend:
        warning_table = make_full_width_black_table(document, 1, 1)
        warning_cell = warning_table.cell(0, 0)
        add_text_to_cell(
            warning_cell,
            "No se encontró la imagen de leyenda en public/leyenda_alergenos.png",
            size=9,
            italic=True,
            color="FFFFFF"
        )

    if include_footer_text:
        footer_table = make_full_width_black_table(document, 1, 1)
        footer_cell = footer_table.cell(0, 0)
        add_text_to_cell(
            footer_cell,
            "Informamos de acuerdo con el Reglamento de la U.E 1169/2011, que nuestros productos contienen o pueden contener los siguientes alérgenos.",
            size=9,
            italic=True,
            color="FFFFFF",
            align=WD_ALIGN_PARAGRAPH.LEFT
        )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =========================================================
# INTERFAZ STREAMLIT
# =========================================================
st.set_page_config(
    page_title=APP_TITLE,
    page_icon="📋",
    layout="wide",
)

st.title("📋 Carta de Alérgenos")
st.caption("Generador de carta editable. La plantilla Word texto limpio elimina los iconos por plato y mantiene la leyenda inferior desde la carpeta public.")

with st.sidebar:
    st.header("Configuración")
    restaurant_name = st.text_input("Título principal", value="MENÚ")

    template = st.selectbox(
        "Plantilla",
        [TEMPLATE_WORD_TEXTO_LIMPIO],
        index=0,
        help="Esta versión exporta texto y precios sin iconos por plato."
    )

    st.divider()
    st.subheader("Leyenda inferior")
    include_footer_legend = st.checkbox("Mantener imagen de leyenda inferior", value=True)
    include_footer_text = st.checkbox("Añadir texto legal debajo", value=True)
    legend_width = st.slider("Ancho de la leyenda en Word", min_value=10.0, max_value=18.0, value=17.2, step=0.1)

    legend_path = find_legend_image()
    if legend_path:
        st.success(f"Leyenda detectada: {legend_path.as_posix()}")
    else:
        st.warning("No se detectó la leyenda. Coloca la imagen en public/leyenda_alergenos.png")

st.subheader("Contenido del menú")
st.write("Formato: `Plato | Precio | alérgenos opcionales`. En **Word texto limpio**, la tercera columna se ignora para que no salgan iconos por plato.")

raw_text = st.text_area(
    "Pega o edita aquí la carta",
    value=DEFAULT_MENU_TEXT,
    height=520,
)

sections = parse_menu_text(raw_text)

col1, col2 = st.columns([1.2, 0.8])

with col1:
    st.subheader("Vista rápida")
    for section in sections:
        if section["title"]:
            st.markdown(f"### {section['title']}")
        for item in section["items"]:
            left, right = st.columns([0.78, 0.22])
            left.write(f"**{item['name']}**")
            right.write(f"**{item['price']}**")

    st.info("En esta plantilla no se muestran iconos por plato. Solo texto, precios y leyenda inferior.")

with col2:
    st.subheader("Exportar")
    if template == TEMPLATE_WORD_TEXTO_LIMPIO:
        docx_bytes = build_word_texto_limpio_docx(
            restaurant_name=restaurant_name,
            sections=sections,
            include_footer_legend=include_footer_legend,
            include_footer_text=include_footer_text,
            legend_width_cm=legend_width,
        )

        filename = f"{clean_filename(restaurant_name)}_word_texto_limpio.docx"

        st.download_button(
            label="⬇️ Descargar Word texto limpio",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.markdown("""
**Qué hace esta versión:**

- Elimina los iconos de alérgenos junto a cada plato.
- Alinea los precios más a la derecha.
- Mantiene la leyenda inferior desde `public`.
- No borra ni reemplaza la imagen existente.
- Genera un Word editable.
""")
